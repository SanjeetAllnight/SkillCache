"""
generate_exp12.py  –  Generates Experiment 12 (.docx) for SEPM project
Source Code Documentation for SkillCache
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── Config ──────────────────────────────────────────────────────────────────

BODY_FONT    = "Times New Roman"
CODE_FONT    = "Courier New"
BODY_SIZE    = Pt(20)
CODE_SIZE    = Pt(18)
HEADING1_SIZE = Pt(28)
HEADING2_SIZE = Pt(24)
HEADING3_SIZE = Pt(20)

doc = Document()

# ─── Page Setup ──────────────────────────────────────────────────────────────

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)


# ─── Style Defaults ─────────────────────────────────────────────────────────

style = doc.styles['Normal']
font  = style.font
font.name = BODY_FONT
font.size = BODY_SIZE
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(3)


# ─── Helper Functions ────────────────────────────────────────────────────────

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = BODY_FONT
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = HEADING1_SIZE
        elif level == 2:
            run.font.size = HEADING2_SIZE
        else:
            run.font.size = HEADING3_SIZE
    return h


def add_body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
    return p


def add_code_block(code_text, label=None):
    """Add a monospace code block with a label."""
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
        run.bold = True
    
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run(line)
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
        run.font.color.rgb = RGBColor(20, 20, 20)


def add_function_doc(name, signature, description):
    p = doc.add_paragraph()
    run = p.add_run(f"• {name}")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"  Signature: {signature}")
    run2.font.name = CODE_FONT
    run2.font.size = CODE_SIZE
    
    p3 = doc.add_paragraph()
    run3 = p3.add_run(f"  {description}")
    run3.font.name = BODY_FONT
    run3.font.size = BODY_SIZE


def set_cell_font(cell, text, font_name=BODY_FONT, size=BODY_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = size
    run.bold = bold


def shade_cells(row, color="D9E2F3"):
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)


# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Shree L. R. Tiwari College of Engineering")
run.font.name = BODY_FONT
run.font.size = Pt(26)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Department of Computer Engineering")
run.font.name = BODY_FONT
run.font.size = Pt(22)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Software Engineering & Project Management (SEPM)")
run.font.name = BODY_FONT
run.font.size = Pt(24)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Experiment No. 12")
run.font.name = BODY_FONT
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Source Code Documentation")
run.font.name = BODY_FONT
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Project: SkillCache – Peer-to-Peer Skill Exchange Platform")
run.font.name = BODY_FONT
run.font.size = Pt(20)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Team Table
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Team Members")
run.font.name = BODY_FONT
run.font.size = Pt(22)
run.bold = True

doc.add_paragraph()

team = [
    ("1", "Sanjeet Ghadi", "56", "9284454860", "sanjeetghadi09@gmail.com"),
    ("2", "Samiksha Mangji", "54", "9699794627", "mangjisamiksha011@gmail.com"),
    ("3", "Prasham Satarkar", "44", "7263010576", "prashamsatarkar@gmail.com"),
    ("4", "Santosh Naik", "57", "8208849682", "naiksantosh663@gmail.com"),
    ("5", "Sakshi Ghadi", "51", "9545248202", "shghadi06@gmail.com"),
]

table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

headers = ["Sr No.", "Name", "Roll No.", "Mobile", "Email"]
for i, header in enumerate(headers):
    set_cell_font(table.rows[0].cells[i], header, bold=True, size=Pt(18))
shade_cells(table.rows[0], "003366")
for cell in table.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

for member in team:
    row = table.add_row()
    for i, val in enumerate(member):
        set_cell_font(row.cells[i], val, size=Pt(18))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  INDEX PAGE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("Index", 1)

index_table = doc.add_table(rows=1, cols=3)
index_table.alignment = WD_TABLE_ALIGNMENT.CENTER
index_table.style = 'Table Grid'

idx_headers = ["Sr No.", "Topic", "Page No."]
for i, h in enumerate(idx_headers):
    set_cell_font(index_table.rows[0].cells[i], h, bold=True, size=Pt(18))
shade_cells(index_table.rows[0], "003366")
for cell in index_table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

topics = [
    ("1", "Introduction"),
    ("2", "Project Overview"),
    ("3", "Folder Structure"),
    ("4", "Module-wise Code Documentation"),
    ("5", "Key Algorithms in Code"),
    ("6", "Database Schema"),
    ("7", "Files Excluded from Submission"),
    ("8", "Conclusion"),
    ("9", "References"),
]

for idx, (num, topic) in enumerate(topics):
    row = index_table.add_row()
    set_cell_font(row.cells[0], num, size=Pt(18))
    set_cell_font(row.cells[1], topic, size=Pt(18), align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(row.cells[2], "-", size=Pt(18))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("1. Introduction", 1)

add_body(
    "Source code documentation is the systematic practice of annotating, describing, "
    "and cataloguing the source code of a software project so that developers, reviewers, "
    "and future maintainers can understand the architecture, data flow, and implementation "
    "decisions without requiring verbal explanations from the original authors. It encompasses "
    "inline comments, function-level docstrings, module-level summaries, type annotations, "
    "and external reference documents that together form a comprehensive knowledge base "
    "around the codebase."
)

add_body(
    "The purpose of source code documentation in the context of Software Engineering and "
    "Project Management (SEPM) is threefold: first, it facilitates knowledge transfer among "
    "team members and reduces the bus factor; second, it serves as a verification artifact "
    "during code reviews and quality audits; and third, it provides a traceable link between "
    "software requirements and their concrete implementations, which is essential for "
    "maintenance, debugging, and future evolution of the system."
)

add_body(
    "This document presents the complete source code documentation for the SkillCache "
    "platform — a peer-to-peer skill exchange web application. It covers the project's "
    "folder structure, module-wise breakdown, key algorithms, database schema, and the "
    "rationale behind design decisions embedded in the code."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("2. Project Overview", 1)

add_body(
    "SkillCache is a full-stack peer-to-peer skill exchange platform that enables users to "
    "discover mentors, book one-on-one video sessions, exchange knowledge resources, and "
    "rate their mentorship experiences. The platform is built as a modern single-page "
    "application with real-time capabilities."
)

add_heading_styled("2.1 Technology Stack", 2)

tech_items = [
    ("Framework", "Next.js 16 (App Router) with React 19"),
    ("Language", "TypeScript 6.0"),
    ("Backend / BaaS", "Firebase (Firestore, Auth, Storage)"),
    ("Real-time Video", "Native WebRTC (RTCPeerConnection) with Firestore-backed signaling"),
    ("Styling", "Tailwind CSS 3.4 with a custom Material Design 3 token system"),
    ("Package Manager", "npm with package-lock.json"),
    ("Linting", "ESLint 9 with Next.js config"),
    ("State Management", "React Context API (AuthProvider)"),
    ("Build Tool", "Next.js built-in (Turbopack/Webpack)"),
]

for label, value in tech_items:
    p = doc.add_paragraph()
    run = p.add_run(f"• {label}: ")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run2 = p.add_run(value)
    run2.font.name = BODY_FONT
    run2.font.size = BODY_SIZE

add_heading_styled("2.2 Codebase Metrics", 2)

add_body("Total Files: 52 source files (excluding node_modules, .next, stitch/, .env, .git)")
add_body("Total Lines of Code: ~9,424 lines across all source files")
add_body("Primary Language: TypeScript / TSX (100% of application code)")
add_body("Firestore Collections: users, sessions, calls, resources, resourceLikes, resourceBookmarks, workshops")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. FOLDER STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("3. Folder Structure", 1)

add_body("The following is the complete file tree of the SkillCache project, excluding node_modules, .next, stitch/, .env, and .git directories. Each file is annotated with a one-line description.")

folder_tree = """SkillCache/
├── app/
│   ├── globals.css                    # Global CSS with MD3 design tokens and utility classes
│   ├── layout.tsx                     # Root layout with metadata, fonts, and providers
│   ├── loading.tsx                    # Root-level loading skeleton component
│   ├── page.tsx                       # Landing page / marketing homepage
│   ├── auth/
│   │   └── page.tsx                   # Auth route wrapper (redirects to auth view)
│   ├── complete-profile/
│   │   └── page.tsx                   # Post-signup onboarding profile completion form
│   ├── session/
│   │   ├── page.tsx                   # Session route index (redirect handler)
│   │   └── [sessionId]/
│   │       └── page.tsx               # Live video call session page with WebRTC
│   └── (shell)/
│       ├── layout.tsx                 # Shell layout with sidebar, navbar, mobile nav
│       ├── loading.tsx                # Shell-level loading skeleton
│       ├── dashboard/
│       │   └── page.tsx               # User dashboard with sessions, stats, mentors
│       ├── discover/
│       │   └── page.tsx               # Community discovery page with trending skills
│       ├── mentors/
│       │   └── page.tsx               # Mentor search and listing page
│       ├── profile/
│       │   └── page.tsx               # User/mentor profile page with skills and reviews
│       ├── repository/
│       │   ├── page.tsx               # Knowledge repository listing and search
│       │   └── [id]/
│       │       └── page.tsx           # Single resource detail page
│       └── sessions/
│           ├── page.tsx               # Sessions listing (upcoming, past, pending)
│           └── [id]/
│               └── page.tsx           # Session detail with lifecycle management
├── components/
│   ├── auth/
│   │   └── auth-page-view.tsx         # Login/signup form with Google OAuth
│   ├── cards/
│   │   ├── mentor-card.tsx            # Reusable mentor profile card component
│   │   ├── resource-card.tsx          # Legacy resource card (deprecated)
│   │   └── session-card.tsx           # Multi-variant session card component
│   ├── providers/
│   │   └── auth-provider.tsx          # React context for Firebase auth state
│   ├── repository/
│   │   ├── knowledge-resource-card.tsx # Repository resource card with actions
│   │   ├── repository-empty-state.tsx # Empty state placeholder for repository
│   │   └── resource-composer.tsx      # Multi-type resource creation form
│   ├── reviews/
│   │   ├── review-modal.tsx           # Star rating review submission modal
│   │   ├── reviews-section.tsx        # Reviews list with real-time updates
│   │   └── star-rating.tsx            # Interactive/display star rating widget
│   ├── sessions/
│   │   ├── session-lifecycle-card.tsx  # Session lifecycle state machine UI
│   │   ├── session-request-modal.tsx  # Session booking request form modal
│   │   └── session-status-badge.tsx   # Colored status badge component
│   ├── shell/
│   │   ├── mobile-nav.tsx             # Bottom navigation bar for mobile
│   │   ├── navbar.tsx                 # Top navigation bar with search
│   │   └── sidebar.tsx               # Desktop sidebar navigation
│   ├── skills/
│   │   ├── ai-import-modal.tsx        # AI-powered skill import from GitHub/LinkedIn
│   │   ├── skill-card.tsx             # Individual skill display card
│   │   ├── skill-importer.tsx         # Skill import trigger button
│   │   ├── skill-modal.tsx            # Add/edit skill form modal
│   │   └── toast.tsx                  # Toast notification component
│   └── ui/
│       ├── button.tsx                 # Reusable button with variants
│       ├── icon.tsx                   # Material Symbols icon wrapper
│       ├── skeleton.tsx               # Loading skeleton placeholder
│       └── tag.tsx                    # Tag/chip display component
├── hooks/
│   └── useWebRTC.ts                   # WebRTC peer connection React hook
├── lib/
│   ├── ai-skills.ts                   # Simulated AI profile scan for skills
│   ├── auth.ts                        # Auth cookie/storage utilities
│   ├── callSignaling.ts              # Firestore-backed WebRTC signaling layer
│   ├── firebase.ts                    # Firebase app initialization
│   ├── firebaseServices.ts           # Core Firestore CRUD for all collections
│   ├── mockUser.ts                    # User type definitions and display helpers
│   ├── repository.ts                  # Knowledge repository CRUD with file uploads
│   ├── seedData.ts                    # Demo data seeding for development
│   ├── shell-config.ts               # Navigation and layout configuration
│   ├── static-assets.ts              # Static asset URL constants
│   ├── utils.ts                       # General utility functions (cn helper)
│   ├── view-models.ts                # Data transformation for UI components
│   └── webrtc.ts                      # Legacy simple-peer WebRTC (deprecated)
├── types/                             # TypeScript type declarations (empty)
├── firestore.rules                    # Firestore security rules (245 lines)
├── storage.rules                      # Firebase Storage security rules
├── middleware.ts                       # Next.js edge middleware for auth guards
├── next.config.ts                     # Next.js configuration
├── next-env.d.ts                      # Next.js TypeScript declarations
├── package.json                       # Project dependencies and scripts
├── postcss.config.js                  # PostCSS configuration for Tailwind
├── tailwind.config.ts                 # Tailwind CSS configuration with MD3 tokens
├── tsconfig.json                      # TypeScript compiler configuration
└── log.txt                            # Development log"""

add_code_block(folder_tree)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. MODULE-WISE CODE DOCUMENTATION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("4. Module-wise Code Documentation", 1)

add_body(
    "The SkillCache platform is logically organised into seven functional modules. "
    "Each module encapsulates a distinct domain of the application. This section "
    "documents the files, key functions, and representative code for every module."
)

# ─── M1: Auth & Profile ──────────────────────────────────────────────────────

add_heading_styled("4.1 Module 1: Auth & Profile", 2)

add_body("This module handles user authentication (email/password and Google OAuth), session cookie management, profile creation, onboarding, and route protection via Next.js middleware.")

add_heading_styled("Files", 3)
add_body("• lib/auth.ts — Auth cookie constants, protected path detection, storage cleanup")
add_body("• lib/firebase.ts — Firebase app initialisation (Auth, Firestore, Storage)")
add_body("• lib/mockUser.ts — BackendUser type definitions, toDisplayUser mapping")
add_body("• components/providers/auth-provider.tsx — AuthContext React provider with login/signup/logout")
add_body("• components/auth/auth-page-view.tsx — Login/signup UI with form validation")
add_body("• app/complete-profile/page.tsx — Post-signup onboarding form")
add_body("• middleware.ts — Next.js edge middleware for auth-gated routes")

add_heading_styled("Key Functions", 3)

add_function_doc(
    "createUserProfile",
    "async (uid: string, data: {name, email}): Promise<void>",
    "Creates the initial Firestore user document on first sign-up with merge semantics."
)
add_function_doc(
    "getUserProfile",
    "async (uid: string): Promise<FirestoreUser | null>",
    "Returns the raw Firestore document for a user, used for onboarding gate decisions."
)
add_function_doc(
    "clearAllAuthStorage",
    "async (): Promise<void>",
    "Nuclear-clears all Firebase auth artefacts from localStorage, sessionStorage, and IndexedDB."
)
add_function_doc(
    "middleware",
    "(request: NextRequest): NextResponse",
    "Edge middleware that redirects unauthenticated users from protected routes to /auth."
)
add_function_doc(
    "isProtectedPath",
    "(pathname: string): boolean",
    "Checks if a given pathname is under a protected route prefix."
)

add_heading_styled("Core Logic — Auth Provider (auth-provider.tsx)", 3)

auth_code = '''export function AuthProvider({ initialIsLoggedIn, children }) {
  const [isAuthReady, setIsAuthReady] = useState(false);
  const [isLoggedIn, setIsLoggedIn] = useState(initialIsLoggedIn);
  const [user, setUser] = useState<BackendUser | null>(null);

  useEffect(() => {
    const unsubscribe = onAuthStateChanged(auth, async (firebaseUser) => {
      if (firebaseUser) {
        await syncUserFromFirestore(firebaseUser.uid, firebaseUser.email);
        void updateLastLoginAt(firebaseUser.uid);
        setIsLoggedIn(true);
        setBrowserAuthCookie(true);

        // Profile-completion gate
        const profile = await getUserProfile(firebaseUser.uid);
        if (!profile || profile.profileComplete === false) {
          routerRef.current.replace("/complete-profile");
        }
      } else {
        setUser(null);
        setIsLoggedIn(false);
        setBrowserAuthCookie(false);
        if (isProtectedPath(pathnameRef.current)) {
          routerRef.current.replace(`/auth?next=${pathnameRef.current}`);
        }
      }
      setIsAuthReady(true);
    });
    return () => unsubscribe();
  }, [syncUserFromFirestore]);

  const logout = useCallback(async () => {
    await clearAllAuthStorage();
    setUser(null);
    setIsLoggedIn(false);
    await signOut(auth);
    window.location.href = "/auth";
  }, []);

  return <AuthContext.Provider value={value}>{children}</AuthContext.Provider>;
}'''
add_code_block(auth_code)

doc.add_page_break()

# ─── M2: Skill Management ────────────────────────────────────────────────────

add_heading_styled("4.2 Module 2: Skill Management", 2)

add_body("This module manages the CRUD lifecycle of user skills stored as Firestore sub-collections under each user document. It also includes AI-powered skill import simulation and synchronises skill arrays on the parent user document for backward compatibility with the mentor search.")

add_heading_styled("Files", 3)
add_body("• lib/firebaseServices.ts (skill section, lines 906–989) — Skill CRUD and real-time listeners")
add_body("• lib/ai-skills.ts — Simulated AI profile scan returning extracted skills")
add_body("• components/skills/skill-modal.tsx — Add/edit skill form with category and level")
add_body("• components/skills/skill-card.tsx — Individual skill display card")
add_body("• components/skills/ai-import-modal.tsx — AI import modal for GitHub/LinkedIn scan")
add_body("• components/skills/skill-importer.tsx — Trigger button for AI import")

add_heading_styled("Key Functions", 3)

add_function_doc(
    "addSkill",
    "async (uid: string, data: Omit<FirestoreSkill, 'createdAt'|'updatedAt'>): Promise<ApiSkill>",
    "Adds a new skill document to the user's skills sub-collection and syncs the flat arrays."
)
add_function_doc(
    "updateSkill",
    "async (uid: string, skillId: string, data: Partial<FirestoreSkill>): Promise<void>",
    "Updates an existing skill document in place and re-syncs the flat skill arrays."
)
add_function_doc(
    "deleteSkill",
    "async (uid: string, skillId: string): Promise<void>",
    "Deletes a skill by its document ID and re-syncs the flat arrays."
)
add_function_doc(
    "syncSkillArrays",
    "async (uid: string): Promise<void>",
    "Reads all skills from the sub-collection and writes skillsOffered/skillsWanted arrays to the user doc."
)
add_function_doc(
    "listenSkillsForUser",
    "(uid: string, onChange: (skills: ApiSkill[]) => void): Unsubscribe",
    "Subscribes to real-time skill updates via Firestore onSnapshot."
)
add_function_doc(
    "simulateAiProfileScan",
    "async (source: 'github' | 'linkedin'): Promise<AiProfileSuggestions>",
    "Simulates AI extraction of skills with confidence scores and suggested roles."
)

add_heading_styled("Core Logic — Skill Sync and Migration (firebaseServices.ts)", 3)

skill_code = '''async function syncSkillArrays(uid: string): Promise<void> {
  const snap = await getDocs(skillsColRef(uid));
  const skills = snap.docs.map((d) => d.data() as FirestoreSkill);
  const offered = skills.filter((s) => s.type === "teaching").map((s) => s.name);
  const wanted  = skills.filter((s) => s.type === "learning").map((s) => s.name);
  await updateDoc(doc(db, "users", uid), {
    skillsOffered: offered,
    skillsWanted:  wanted,
    profileComplete: true,
  });
}

export async function migrateSkillsFromLegacy(uid: string): Promise<boolean> {
  const colRef = skillsColRef(uid);
  const existing = await getDocs(colRef);
  if (!existing.empty) return false;

  const userSnap = await getDoc(doc(db, "users", uid));
  if (!userSnap.exists()) return false;

  const data = userSnap.data() as FirestoreUser;
  const offered: string[] = data.skillsOffered ?? [];
  const wanted: string[]  = data.skillsWanted  ?? [];

  if (offered.length === 0 && wanted.length === 0) return false;

  const seen = new Set<string>();
  const writes: Promise<unknown>[] = [];

  for (const name of offered) {
    const key = name.toLowerCase().trim();
    if (!key || seen.has(key)) continue;
    seen.add(key);
    writes.push(addDoc(colRef, {
      name: name.trim(), category: "Other", tags: [],
      type: "teaching", description: "", level: "intermediate",
      createdAt: serverTimestamp(), updatedAt: serverTimestamp(),
    }));
  }
  await Promise.all(writes);
  return true;
}'''
add_code_block(skill_code)

doc.add_page_break()

# ─── M3: Matching Engine ─────────────────────────────────────────────────────

add_heading_styled("4.3 Module 3: Matching Engine", 2)

add_body("The matching engine is responsible for connecting learners with appropriate mentors. It filters users by skill offerings, ranks them by relevance, and provides personalised recommendations based on the learner's skill wants. The discover page aggregates trending skills across the community.")

add_heading_styled("Files", 3)
add_body("• lib/firebaseServices.ts (getMentors, lines 351–373) — Mentor filtering by skill")
add_body("• lib/view-models.ts — toMentorCardData mapping for UI display")
add_body("• app/(shell)/discover/page.tsx — Community discovery with trending skills and recommendations")
add_body("• app/(shell)/mentors/page.tsx — Mentor search and listing page")
add_body("• components/cards/mentor-card.tsx — Mentor profile card component")

add_heading_styled("Key Functions", 3)

add_function_doc(
    "getMentors",
    "async (skillFilter?: string, excludeUid?: string): Promise<BackendUser[]>",
    "Fetches users who offer at least one skill, with optional case-insensitive substring filtering."
)
add_function_doc(
    "calculateTrendingSkills",
    "(users: BackendUser[]): TrendingSkill[]",
    "Aggregates skill demand/supply across all users and ranks by composite score."
)
add_function_doc(
    "toMentorCardData",
    "(mentors: BackendUser[]): MentorCardData[]",
    "Maps raw BackendUser objects to the MentorCard display format."
)

add_heading_styled("Core Logic — Mentor Matching & Trending Skills", 3)

matching_code = '''export async function getMentors(
  skillFilter?: string,
  excludeUid?: string
): Promise<BackendUser[]> {
  const users = await getUsers();

  return users.filter((user) => {
    if (excludeUid && user._id === excludeUid) return false;
    const hasSkills = user.skillsOffered && user.skillsOffered.length > 0;
    if (!hasSkills) return false;

    if (skillFilter) {
      const q = skillFilter.toLowerCase();
      return user.skillsOffered!.some((s) => s.toLowerCase().includes(q));
    }
    return true;
  });
}

function calculateTrendingSkills(users: BackendUser[]) {
  const skillStats: Record<string, { learners: number; mentors: number; trend: number }> = {};

  users.forEach((u) => {
    u.skillsWanted?.forEach((s) => {
      if (!skillStats[s]) skillStats[s] = { learners: 0, mentors: 0, trend: Math.floor(Math.random() * 20) + 5 };
      skillStats[s].learners++;
    });
    u.skillsOffered?.forEach((s) => {
      if (!skillStats[s]) skillStats[s] = { learners: 0, mentors: 0, trend: Math.floor(Math.random() * 20) + 5 };
      skillStats[s].mentors++;
    });
  });

  return Object.entries(skillStats)
    .map(([name, stats]) => ({ name, ...stats, score: stats.learners + stats.mentors * 2 }))
    .sort((a, b) => b.score - a.score)
    .slice(0, 4);
}'''
add_code_block(matching_code)

doc.add_page_break()

# ─── M4: Booking / Session System ────────────────────────────────────────────

add_heading_styled("4.4 Module 4: Booking / Session System", 2)

add_body("This module manages the complete session lifecycle from booking request through completion or cancellation. It includes real-time session listeners, status transitions, and time-based automatic promotion and expiry of sessions.")

add_heading_styled("Files", 3)
add_body("• lib/firebaseServices.ts (session section, lines 377–671) — Session CRUD, lifecycle, listeners")
add_body("• components/sessions/session-request-modal.tsx — Session booking form")
add_body("• components/sessions/session-lifecycle-card.tsx — Session state machine UI")
add_body("• components/sessions/session-status-badge.tsx — Status badge component")
add_body("• app/(shell)/sessions/page.tsx — Sessions listing")
add_body("• app/(shell)/sessions/[id]/page.tsx — Session detail with lifecycle controls")

add_heading_styled("Key Functions", 3)

add_function_doc(
    "requestSession",
    "async (data: {...}): Promise<string>",
    "Creates a session with 'pending' status and marks the requester as the learner."
)
add_function_doc(
    "acceptSession",
    "async (sessionId: string): Promise<void>",
    "Transitions a session from 'pending' to 'accepted' and stamps acceptedAt."
)
add_function_doc(
    "startSession",
    "async (sessionId: string, userId: string): Promise<void>",
    "Transitions a session to 'live' and sets callStatus to 'started'."
)
add_function_doc(
    "completeSession",
    "async (sessionId: string, userId: string): Promise<void>",
    "Marks a session as 'completed' and resets callStatus to 'idle'."
)
add_function_doc(
    "syncSessionLifecycle",
    "async (session: ApiSession, now?: number): Promise<void>",
    "Automatically promotes 'accepted' sessions to 'upcoming' or marks them as 'missed' based on time windows."
)
add_function_doc(
    "canJoinSession",
    "(session: ApiSession, now?: number): boolean",
    "Returns true if a session is within the joinable time window (15min before to 2hr after)."
)
add_function_doc(
    "listenSessionsForUser",
    "(uid: string, onChange: Function): Unsubscribe",
    "Attaches dual Firestore listeners (mentor + learner queries) with deduplication."
)

add_heading_styled("Core Logic — Session Lifecycle State Machine", 3)

session_code = '''export async function syncSessionLifecycle(session: ApiSession, now = Date.now()): Promise<void> {
  const startsAt = new Date(session.date).getTime();
  if (!Number.isFinite(startsAt)) return;

  const shouldBecomeUpcoming =
    session.status === "accepted" &&
    now >= startsAt - SESSION_JOIN_WINDOW_MS &&
    now <= startsAt + SESSION_MISS_GRACE_MS;

  if (shouldBecomeUpcoming) {
    await promoteSessionToUpcoming(session._id);
    return;
  }

  const shouldBeMissed =
    (session.status === "accepted" || session.status === "upcoming") &&
    now > startsAt + SESSION_MISS_GRACE_MS;

  if (shouldBeMissed) {
    await markSessionMissed(session._id);
  }
}

export function canJoinSession(session: ApiSession, now = Date.now()) {
  const startsAt = new Date(session.date).getTime();
  if (!Number.isFinite(startsAt)) return false;
  if (session.status === "live") return true;
  if (session.status !== "accepted" && session.status !== "upcoming") return false;
  return now >= startsAt - SESSION_JOIN_WINDOW_MS && now <= startsAt + SESSION_MISS_GRACE_MS;
}

export const SESSION_JOIN_WINDOW_MS = 15 * 60 * 1000;   // 15 minutes
export const SESSION_MISS_GRACE_MS  = 2 * 60 * 60 * 1000; // 2 hours'''
add_code_block(session_code)

doc.add_page_break()

# ─── M5: Repository ──────────────────────────────────────────────────────────

add_heading_styled("4.5 Module 5: Repository", 2)

add_body("The repository module implements a full-featured knowledge base where users can share, discover, and manage learning resources. It supports multiple resource types (PDF, markdown, code, links, images), file uploads with retry logic, likes/bookmarks with transactional counters, and paginated listing with local text search.")

add_heading_styled("Files", 3)
add_body("• lib/repository.ts — Complete repository CRUD with file upload, transactions, pagination (830 lines)")
add_body("• components/repository/resource-composer.tsx — Multi-type resource creation form")
add_body("• components/repository/knowledge-resource-card.tsx — Resource card with like/bookmark actions")
add_body("• components/repository/repository-empty-state.tsx — Empty state placeholder")
add_body("• app/(shell)/repository/page.tsx — Repository listing with search and tag filtering")
add_body("• app/(shell)/repository/[id]/page.tsx — Single resource detail page")

add_heading_styled("Key Functions", 3)

add_function_doc(
    "createKnowledgeResource",
    "async (input: CreateResourceInput, onProgress?): Promise<KnowledgeResource>",
    "Validates input, writes Firestore doc, uploads file with retry, and returns optimistic resource."
)
add_function_doc(
    "uploadFileWithRetry",
    "async (path, file, metadata, onProgress?, maxRetries?, timeout?): Promise<string>",
    "Uploads a file to Firebase Storage with exponential backoff retry and timeout protection."
)
add_function_doc(
    "toggleResourceLike",
    "async (resourceId: string, userId: string): Promise<boolean>",
    "Atomically toggles a like using Firestore transactions with counter increment/decrement."
)
add_function_doc(
    "toggleResourceBookmark",
    "async (resourceId: string, userId: string): Promise<boolean>",
    "Atomically toggles a bookmark using Firestore transactions."
)
add_function_doc(
    "listRepositoryResources",
    "async (options: {...}): Promise<ResourcePage>",
    "Paginated listing of resources with scope filtering (community/mine) and cursor-based pagination."
)
add_function_doc(
    "filterResourcesLocally",
    "(resources[], search, activeTag?): KnowledgeResource[]",
    "Client-side text search across title, description, tags, and uploader name."
)
add_function_doc(
    "deleteKnowledgeResource",
    "async (resource, currentUserId): Promise<void>",
    "Deletes the resource document, storage file, and all associated likes/bookmarks in a batch."
)

add_heading_styled("Core Logic — Transactional Like Toggle (repository.ts)", 3)

repo_code = '''export async function toggleResourceLike(resourceId: string, userId: string) {
  const resourceRef = doc(resourcesCollection, resourceId);
  const likeRef = doc(likesCollection, `${resourceId}_${userId}`);

  return runTransaction(db, async (transaction) => {
    const [resourceSnap, likeSnap] = await Promise.all([
      transaction.get(resourceRef),
      transaction.get(likeRef),
    ]);

    if (!resourceSnap.exists()) throw new Error("This resource no longer exists.");

    if (likeSnap.exists()) {
      transaction.delete(likeRef);
      transaction.update(resourceRef, {
        likesCount: increment(-1),
        updatedAt: serverTimestamp(),
      });
      return false;
    }

    transaction.set(likeRef, {
      resourceId, userId,
      createdAt: serverTimestamp(),
    });
    transaction.update(resourceRef, {
      likesCount: increment(1),
      updatedAt: serverTimestamp(),
    });
    return true;
  });
}'''
add_code_block(repo_code)

doc.add_page_break()

# ─── M6: Rating System ───────────────────────────────────────────────────────

add_heading_styled("4.6 Module 6: Rating System", 2)

add_body("The rating system allows learners to submit star ratings (1–5) and text reviews for mentors after sessions. It enforces business rules (no self-review, no duplicates, character limits), uses an incremental average algorithm to maintain the mentor's aggregate rating, and provides real-time review streams.")

add_heading_styled("Files", 3)
add_body("• lib/firebaseServices.ts (review section, lines 1063–1231) — Review CRUD and rating calculations")
add_body("• components/reviews/review-modal.tsx — Star rating review submission modal")
add_body("• components/reviews/reviews-section.tsx — Real-time reviews list with rating summary")
add_body("• components/reviews/star-rating.tsx — Interactive/display star rating widget")

add_heading_styled("Key Functions", 3)

add_function_doc(
    "submitReview",
    "async (mentorUid, reviewer, input): Promise<ApiReview>",
    "Validates, writes review doc, and atomically recalculates the mentor's incremental average rating."
)
add_function_doc(
    "hasUserReviewedMentor",
    "async (mentorUid, reviewerUid): Promise<boolean>",
    "Checks for existing review to prevent duplicates."
)
add_function_doc(
    "getReviewsForMentor",
    "async (mentorUid): Promise<ApiReview[]>",
    "Fetches all reviews for a mentor, sorted newest-first."
)
add_function_doc(
    "listenReviewsForMentor",
    "(mentorUid, onChange): () => void",
    "Real-time Firestore listener for a mentor's reviews sub-collection."
)
add_function_doc(
    "getMentorRatingMeta",
    "async (mentorUid): Promise<MentorRatingMeta>",
    "Reads averageRating + totalReviews from the mentor's user document."
)

add_heading_styled("Core Logic — Incremental Average Rating (firebaseServices.ts)", 3)

rating_code = '''export async function submitReview(
  mentorUid: string,
  reviewer: { uid: string; name: string },
  input: { rating: number; reviewText: string; sessionId?: string | null },
): Promise<ApiReview> {
  if (mentorUid === reviewer.uid) throw new Error("You cannot review yourself.");
  if (input.rating < 1 || input.rating > 5) throw new Error("Please select a star rating.");
  if (!input.reviewText.trim() || input.reviewText.trim().length < 5)
    throw new Error("Please write at least 5 characters in your review.");
  if (input.reviewText.length > 300)
    throw new Error("Reviews must be 300 characters or fewer.");

  const alreadyReviewed = await hasUserReviewedMentor(mentorUid, reviewer.uid);
  if (alreadyReviewed) throw new Error("You've already reviewed this mentor.");

  // Write the review document
  const colRef = reviewsColRef(mentorUid);
  const docRef = await addDoc(colRef, {
    reviewerUid: reviewer.uid, reviewerName: reviewer.name,
    rating: input.rating, reviewText: input.reviewText.trim(),
    sessionId: input.sessionId ?? null, createdAt: serverTimestamp(),
  });

  // Atomically recalculate incremental average
  const mentorRef = doc(db, "users", mentorUid);
  const mentorSnap = await getDoc(mentorRef);
  const mentorData = mentorSnap.data() ?? {};
  const prevTotal: number = Number(mentorData.totalReviews ?? 0);
  const prevAvg: number   = Number(mentorData.averageRating ?? 0);
  const newTotal = prevTotal + 1;
  const newAvg = parseFloat(((prevAvg * prevTotal + input.rating) / newTotal).toFixed(2));

  await updateDoc(mentorRef, { averageRating: newAvg, totalReviews: newTotal });

  return { _id: docRef.id, ...input, reviewerUid: reviewer.uid,
    reviewerName: reviewer.name, createdAt: null, createdAtMillis: Date.now() };
}'''
add_code_block(rating_code)

doc.add_page_break()

# ─── M7: Video Session (WebRTC) ──────────────────────────────────────────────

add_heading_styled("4.7 Module 7: Video Session (WebRTC)", 2)

add_body("This module implements real-time peer-to-peer video calling using native WebRTC (RTCPeerConnection) with a Firestore-backed signaling layer. The mentor acts as the initiator (creates the SDP offer), and the learner acts as the receiver (creates the SDP answer). ICE candidates are exchanged through a Firestore sub-collection with automatic self-filtering.")

add_heading_styled("Files", 3)
add_body("• hooks/useWebRTC.ts — React hook wrapping RTCPeerConnection lifecycle (492 lines)")
add_body("• lib/callSignaling.ts — Firestore-backed signaling: offer/answer/ICE exchange (314 lines)")
add_body("• app/session/[sessionId]/page.tsx — Video call UI page")

add_heading_styled("Key Functions", 3)

add_function_doc(
    "useWebRTC",
    "(options: UseWebRTCOptions): UseWebRTCReturn",
    "React hook that manages RTCPeerConnection, media acquisition (3-tier fallback), and call state."
)
add_function_doc(
    "createCall",
    "async (callId, callerUid, offer): Promise<void>",
    "Creates the Firestore call document with SDP offer and marks the caller as initiator."
)
add_function_doc(
    "joinCall",
    "async (callId, calleeUid, answer): Promise<void>",
    "Writes SDP answer to Firestore and adds callee to participants array."
)
add_function_doc(
    "sendSignal",
    "async (callId, senderUid, candidate): Promise<void>",
    "Writes an ICE candidate to the candidates sub-collection for the remote peer."
)
add_function_doc(
    "listenForSignals",
    "(callId, myUid, options): Unsubscribe",
    "Dual Firestore listener: watches call doc for offer/answer and candidates sub-collection for ICE."
)
add_function_doc(
    "endCall",
    "async (callId): Promise<void>",
    "Marks the call as ended in Firestore (best-effort)."
)
add_function_doc(
    "acquireMedia",
    "async (): Promise<MediaStream>",
    "3-tier media fallback: full (video+audio) → audio-only → empty stream. Never throws."
)

add_heading_styled("Core Logic — WebRTC Signaling Handshake (callSignaling.ts)", 3)

webrtc_code = '''export function listenForSignals(
  callId: string, myUid: string, options: SignalListenerOptions
): Unsubscribe {
  const { onOffer, onAnswer, onCandidate, onStatusChange } = options;
  let offerDelivered  = false;
  let answerDelivered = false;
  let lastStatus: CallStatus | null = null;

  // 1. Watch the call document (offer / answer / status)
  const unsubDoc = onSnapshot(callDocRef(callId), (snap) => {
    if (!snap.exists()) return;
    const data = snap.data() as DocumentData;

    // Offer → only deliver to the NON-initiator, only once
    if (onOffer && data.offer && !offerDelivered && data.initiatorUid !== myUid) {
      offerDelivered = true;
      onOffer(data.offer as RTCSessionDescriptionInit);
    }

    // Answer → only deliver to the INITIATOR, only once
    if (onAnswer && data.answer && !answerDelivered && data.initiatorUid === myUid) {
      answerDelivered = true;
      onAnswer(data.answer as RTCSessionDescriptionInit);
    }

    if (onStatusChange && data.status && data.status !== lastStatus) {
      lastStatus = data.status as CallStatus;
      onStatusChange(data.status as CallStatus);
    }
  });

  // 2. Watch the candidates sub-collection
  const unsubCandidates = onSnapshot(candidatesRef(callId), (snap) => {
    snap.docChanges().forEach((change) => {
      if (change.type !== "added") return;
      const data = change.doc.data() as CandidateDocument;
      if (data.uid === myUid) return; // Ignore own candidates
      if (onCandidate && data.candidate) onCandidate(data.candidate);
    });
  });

  return () => { unsubDoc(); unsubCandidates(); };
}'''
add_code_block(webrtc_code)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. KEY ALGORITHMS IN CODE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("5. Key Algorithms in Code", 1)

add_body("This section presents the actual implementation code for the four most algorithmically significant pieces of the SkillCache platform.")

# 5.1 Mentor Ranking
add_heading_styled("5.1 Mentor Ranking / Matching Logic", 2)

add_body("The mentor matching algorithm uses a multi-criteria filter: it excludes the current user, requires at least one offered skill, and supports case-insensitive partial substring matching against the mentor's skill set. The discover page additionally ranks mentors by number of offered skills and computes a composite trending score for skills across the community.")

matching_algo = '''// Mentor filtering with case-insensitive substring match
export async function getMentors(skillFilter?: string, excludeUid?: string): Promise<BackendUser[]> {
  const users = await getUsers();
  return users.filter((user) => {
    if (excludeUid && user._id === excludeUid) return false;
    const hasSkills = user.skillsOffered && user.skillsOffered.length > 0;
    if (!hasSkills) return false;
    if (skillFilter) {
      const q = skillFilter.toLowerCase();
      return user.skillsOffered!.some((s) => s.toLowerCase().includes(q));
    }
    return true;
  });
}

// Recommendation: mentors who offer what the user wants to learn
const userWants = user?.skillsWanted || [];
const recommendedMentors = featuredMentors.filter((m) =>
  m.skillsOffered?.some((skill) => userWants.includes(skill))
);

// Trending skills composite score: demand (learners) + supply*2 (mentors)
score = stats.learners + stats.mentors * 2;'''
add_code_block(matching_algo)

# 5.2 Session Status Transition
add_heading_styled("5.2 Session Status Transition Logic", 2)

add_body("The session lifecycle follows a finite state machine with six states: pending → accepted → upcoming → live → completed, with missed and cancelled as terminal states. Transitions are time-driven (join window, grace period) and user-driven (accept, start, complete, cancel).")

session_algo = '''// Status transitions:
//   pending   → accepted   (mentor accepts)
//   accepted  → upcoming   (auto: within 15min before start)
//   upcoming  → live       (mentor starts call)
//   live      → completed  (either participant ends)
//   accepted/upcoming → missed (auto: 2hr past start)
//   any non-terminal → cancelled (either participant cancels)

export const SESSION_JOIN_WINDOW_MS = 15 * 60 * 1000;    // 15 minutes before
export const SESSION_MISS_GRACE_MS  = 2 * 60 * 60 * 1000; // 2 hours after

export async function syncSessionLifecycle(session: ApiSession, now = Date.now()) {
  const startsAt = new Date(session.date).getTime();
  if (!Number.isFinite(startsAt)) return;

  // Auto-promote accepted → upcoming when within the join window
  if (session.status === "accepted" &&
      now >= startsAt - SESSION_JOIN_WINDOW_MS &&
      now <= startsAt + SESSION_MISS_GRACE_MS) {
    await promoteSessionToUpcoming(session._id);
    return;
  }

  // Auto-expire to missed when past the grace period
  if ((session.status === "accepted" || session.status === "upcoming") &&
      now > startsAt + SESSION_MISS_GRACE_MS) {
    await markSessionMissed(session._id);
  }
}

function normalizeSessionStatus(status?: RawSessionStatus): SessionStatus {
  if (status === "scheduled" || status === "waiting") return "upcoming";
  return status ?? "pending";
}'''
add_code_block(session_algo)

doc.add_page_break()

# 5.3 WebRTC Signaling Handshake
add_heading_styled("5.3 WebRTC Signaling Handshake", 2)

add_body("The WebRTC handshake uses Firestore as the signaling server. The mentor (initiator) creates an SDP offer, writes it to a Firestore document, and listens for the learner's answer. The learner reads the offer, creates an answer, and writes it back. ICE candidates are exchanged through a sub-collection with deduplication guards.")

webrtc_algo = '''// MENTOR (Initiator) — startCall flow:
const offer = await pc.createOffer({
  offerToReceiveAudio: true,
  offerToReceiveVideo: true,
});
await pc.setLocalDescription(offer);
await createCall(callId, myUid, { type: offer.type, sdp: offer.sdp! });

// Listen for answer from learner
const unsub = listenForSignals(callId, myUid, {
  onAnswer: async (answer) => {
    if (hasProcessedAnswerRef.current) return;
    hasProcessedAnswerRef.current = true;
    await pc.setRemoteDescription(new RTCSessionDescription(answer));
    await drainICEQueue(pc);
  },
  onCandidate: handleICE,
});

// LEARNER (Receiver) — joinCall flow:
async function processOffer(offer: RTCSessionDescriptionInit) {
  if (hasProcessedOfferRef.current) return;
  hasProcessedOfferRef.current = true;
  await pc.setRemoteDescription(new RTCSessionDescription(offer));
  await drainICEQueue(pc);

  const answer = await pc.createAnswer();
  await pc.setLocalDescription(answer);
  await firestoreJoinCall(callId, myUid, { type: answer.type, sdp: answer.sdp! });
}

// ICE candidate handling with buffering
async function handleICE(candidate: RTCIceCandidateInit) {
  const pc = pcRef.current;
  if (!pc) return;
  const candidateStr = JSON.stringify(candidate);
  if (addedCandidatesRef.current.has(candidateStr)) return;
  addedCandidatesRef.current.add(candidateStr);

  if (!pc.remoteDescription) {
    iceCandidateQueue.current.push(candidate);  // Buffer until ready
    return;
  }
  await pc.addIceCandidate(new RTCIceCandidate(candidate));
}'''
add_code_block(webrtc_algo)

# 5.4 Rating Average Update
add_heading_styled("5.4 Rating Incremental Average Update", 2)

add_body("The rating system uses an incremental average formula to avoid re-reading all reviews when calculating the new average. This is O(1) per review submission rather than O(n).")

rating_algo = '''// Incremental average formula:
//   newAvg = (prevAvg * prevTotal + newRating) / (prevTotal + 1)

const mentorRef = doc(db, "users", mentorUid);
const mentorSnap = await getDoc(mentorRef);
const mentorData = mentorSnap.data() ?? {};

const prevTotal: number = Number(mentorData.totalReviews ?? 0);
const prevAvg: number   = Number(mentorData.averageRating ?? 0);
const newTotal = prevTotal + 1;
const newAvg = parseFloat(
  ((prevAvg * prevTotal + input.rating) / newTotal).toFixed(2)
);

await updateDoc(mentorRef, {
  averageRating: newAvg,
  totalReviews: newTotal,
});'''
add_code_block(rating_algo)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("6. Database Schema", 1)

add_body("SkillCache uses Google Cloud Firestore as its NoSQL database. The following TypeScript type definitions represent the actual data structures stored in each Firestore collection. These types are defined in lib/firebaseServices.ts and lib/repository.ts and are used throughout the application for type-safe data access.")

add_heading_styled("6.1 Users Collection — FirestoreUser", 2)

user_schema = '''export type FirestoreUser = {
  name: string;
  email: string;
  skillsOffered: string[];
  skillsWanted: string[];
  bio: string;
  rating: number;
  sessionsCompleted: number;
  createdAt?: FieldValue;
  lastLoginAt?: FieldValue;
  profileComplete?: boolean;
  firstLoginCompleted?: boolean;
  isSeeded?: boolean;
};'''
add_code_block(user_schema)

add_heading_styled("6.2 Skills Sub-collection — FirestoreSkill", 2)

skill_schema = '''export type SkillType = "teaching" | "learning";
export type ExperienceLevel = "beginner" | "intermediate" | "advanced" | "expert";

export type FirestoreSkill = {
  name: string;
  category: string;
  tags: string[];
  type: SkillType;
  description: string;
  level: ExperienceLevel;
  createdAt?: FieldValue;
  updatedAt?: FieldValue;
};'''
add_code_block(skill_schema)

add_heading_styled("6.3 Sessions Collection — FirestoreSession", 2)

session_schema = '''export type SessionStatus =
  | "pending" | "accepted" | "upcoming"
  | "live" | "completed" | "cancelled" | "missed";

export type CallStatus = "idle" | "started" | "joined" | "connected";

export type FirestoreSession = {
  title: string;
  description?: string;
  mentorId: string;
  learnerId: string;
  skill: string;
  date: string;
  status: RawSessionStatus;
  requestedBy?: string;
  acceptedAt?: FieldValue;
  startedAt?: FieldValue;
  endedAt?: FieldValue;
  cancelledAt?: FieldValue;
  cancelledBy?: string;
  cancellationReason?: string;
  rescheduleNote?: string;
  mentorJoinedAt?: FieldValue;
  learnerJoinedAt?: FieldValue;
  createdAt?: FieldValue;
  updatedAt?: FieldValue;
  callStatus?: CallStatus;
  isSeeded?: boolean;
};'''
add_code_block(session_schema)

add_heading_styled("6.4 Calls Collection — CallDocument", 2)

call_schema = '''export interface CallDocument {
  offer: RTCSessionDescriptionInit | null;
  answer: RTCSessionDescriptionInit | null;
  participants: string[];
  initiatorUid: string;
  status: CallStatus;        // "ringing" | "active" | "ended"
  createdAt: Timestamp;
  updatedAt: Timestamp;
}

export interface CandidateDocument {
  uid: string;
  candidate: RTCIceCandidateInit;
  createdAt: Timestamp;
}'''
add_code_block(call_schema)

add_heading_styled("6.5 Resources Collection — FirestoreResourceDocument", 2)

resource_schema = '''export type ResourceType = "pdf" | "markdown" | "rich-text" | "code" | "link" | "image" | "file";
export type ResourceVisibility = "public" | "private" | "shared";

export type FirestoreResourceDocument = {
  title: string;
  titleLower: string;
  description: string;
  type: ResourceType;
  uploaderId: string;
  uploaderName: string;
  uploaderAvatar?: string;
  tags: string[];
  tagSlugs: string[];
  sessionId?: string | null;
  sessionParticipantIds?: string[];
  fileUrl?: string;
  storagePath?: string;
  fileName?: string;
  fileSize?: number;
  contentType?: string;
  externalUrl?: string;
  content?: string;
  contentFormat?: "plain" | "markdown" | "rich-text" | "code";
  codeLanguage?: string;
  likesCount: number;
  bookmarksCount: number;
  downloadsCount: number;
  visibility: ResourceVisibility;
  uploadStatus: ResourceUploadStatus;
  ai?: ResourceAiMetadata;
  searchText: string;
  createdAt?: unknown;
  updatedAt?: unknown;
};'''
add_code_block(resource_schema)

add_heading_styled("6.6 Reviews Sub-collection — FirestoreReview", 2)

review_schema = '''export type FirestoreReview = {
  reviewerUid: string;
  reviewerName: string;
  rating: number;          // 1-5
  reviewText: string;
  sessionId?: string | null;
  createdAt?: unknown;
};

export type MentorRatingMeta = {
  averageRating: number;
  totalReviews: number;
};'''
add_code_block(review_schema)

add_heading_styled("6.7 BackendUser (User API Model)", 2)

backend_user_schema = '''export type BackendUser = {
  _id: string;
  name: string;
  email: string;
  skillsOffered?: string[];
  skillsWanted?: string[];
  bio?: string;
  avatar?: string;
  firstLoginCompleted?: boolean;
  isSeeded?: boolean;
};'''
add_code_block(backend_user_schema)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. FILES EXCLUDED FROM SUBMISSION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("7. Files Excluded from Submission", 1)

add_body("The following directories and files are excluded from the project submission. Each exclusion is accompanied by a justification.")

exclusions = [
    ("node_modules/",
     "Contains all third-party npm packages installed via 'npm install'. This directory is "
     "auto-generated from package.json and package-lock.json and typically exceeds 200 MB "
     "in size. It is universally excluded from version control and submission as it can be "
     "recreated deterministically by running 'npm ci' or 'npm install'."),

    (".next/",
     "Contains the Next.js build output including compiled JavaScript bundles, optimised "
     "assets, server-side rendering cache, and webpack artifacts. This directory is "
     "regenerated on every 'npm run build' or 'npm run dev' command and is specific to "
     "the local build environment. Including it would add hundreds of generated files "
     "with no documentary value."),

    ("stitch/",
     "Contains intermediate export files from the Figma-to-code pipeline (Stitch). "
     "These are raw design tokens and component stubs that were used as references "
     "during the initial development phase. They have been fully replaced by the "
     "production code in the app/, components/, and lib/ directories."),

    (".env",
     "Contains environment-specific configuration variables including Firebase API keys, "
     "project IDs, and authentication domain settings. Excluding this file is a mandatory "
     "security practice as it prevents accidental exposure of credentials. A sanitised "
     ".env.example file is provided to document the required variables."),

    (".git/",
     "Contains the Git version control metadata including the full commit history, "
     "branch references, and repository configuration. This directory is maintained "
     "by the Git tool and is excluded from submission as the commit history is not "
     "required for academic evaluation."),
]

for name, reason in exclusions:
    p = doc.add_paragraph()
    run = p.add_run(f"• {name}")
    run.bold = True
    run.font.name = CODE_FONT
    run.font.size = BODY_SIZE
    add_body(reason)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("8. Conclusion", 1)

add_body(
    "This experiment has presented a comprehensive source code documentation of the "
    "SkillCache peer-to-peer skill exchange platform. The documentation covers the "
    "complete architecture across 52 source files organised into seven functional "
    "modules: Authentication and Profile Management, Skill Management, Matching Engine, "
    "Booking and Session System, Knowledge Repository, Rating System, and Real-time "
    "Video Sessions via WebRTC. For each module, the constituent files, key function "
    "signatures, and representative code excerpts have been catalogued. The four most "
    "algorithmically significant implementations — mentor matching logic, session status "
    "transition state machine, WebRTC signaling handshake, and incremental average rating "
    "update — have been documented with their complete source code. Additionally, the "
    "full Firestore database schema has been presented through the TypeScript type "
    "definitions that enforce data integrity across the application. Source code "
    "documentation of this nature is an indispensable practice in professional software "
    "engineering as it ensures maintainability, facilitates onboarding, and provides a "
    "traceable link between requirements and their implementation."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled("9. References", 1)

references = [
    'R. S. Pressman and B. R. Maxim, Software Engineering: A Practitioner\'s Approach, 9th ed. New York, NY: McGraw-Hill Education, 2020.',
    'I. Sommerville, Software Engineering, 10th ed. Harlow, UK: Pearson Education, 2016.',
    'Next.js Documentation, "App Router," Vercel Inc. [Online]. Available: https://nextjs.org/docs/app. [Accessed: Jun. 2026].',
    'Firebase Documentation, "Cloud Firestore," Google LLC. [Online]. Available: https://firebase.google.com/docs/firestore. [Accessed: Jun. 2026].',
    'MDN Web Docs, "WebRTC API," Mozilla Foundation. [Online]. Available: https://developer.mozilla.org/en-US/docs/Web/API/WebRTC_API. [Accessed: Jun. 2026].',
    'TypeScript Documentation, "TypeScript Handbook," Microsoft Corp. [Online]. Available: https://www.typescriptlang.org/docs/handbook/. [Accessed: Jun. 2026].',
    'React Documentation, "React Reference," Meta Platforms Inc. [Online]. Available: https://react.dev/reference/react. [Accessed: Jun. 2026].',
    'R. C. Martin, Clean Code: A Handbook of Agile Software Craftsmanship. Upper Saddle River, NJ: Prentice Hall, 2009.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {ref}")
    run.font.name = BODY_FONT
    run.font.size = Pt(18)


# ═══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Experiment_12_Source_Code_Documentation.docx")
doc.save(output_path)
print(f"\nDone! Document saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
